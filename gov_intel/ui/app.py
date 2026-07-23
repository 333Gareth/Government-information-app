"""The main GOV.UK Policy Intelligence Workstation window.

Threading rules followed throughout this module (this is the main fix
from the original code):

* Background threads NEVER touch Tkinter widgets, ``messagebox``, or
  ``simpledialog`` directly. They only compute data and hand results
  back via ``self.root.after(0, callback)``.
* ``self.log(...)`` is safe to call from any thread -- it schedules the
  actual widget update on the main thread internally.
"""

from __future__ import annotations

import logging
import os
import threading
import webbrowser
from datetime import datetime
from pathlib import Path

import requests
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, simpledialog, ttk

from .. import archive, config
from ..app_state import AppState
from ..gov_api import classify_attachment_url, deep_harvest_gov_uk, sanitize_filename
from ..models import Document
from .data_grid import DataGridViewerWindow
from .pdf_viewer import PDFViewerWidget
from .split_compare import SplitPDFComparatorWindow

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


def _attachment_label(url: str) -> str:
    tail = url.rstrip("/").rsplit("/", 1)[-1].replace("%20", " ")
    return f"{classify_attachment_url(url)}: {tail or url}"


class GovApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("🇬🇧 GOV.UK Policy Intelligence Workstation")
        self.root.geometry("1300x900")

        config.ensure_data_dirs()
        self.state = AppState.load()

        self.active_docs: list[Document] = []
        self.selected_doc: Document | None = None
        self.active_atts: list[str] = []

        self.fav_selected_doc_id: str | None = None
        self.fav_active_atts: list[str] = []
        self.all_fav_attachments: dict[str, list[str]] = {"pdf": [], "link": [], "data": []}

        self.style = ttk.Style()
        self.style.theme_use("clam")

        self.nb = ttk.Notebook(root)
        self.nb.pack(fill="both", expand=True, padx=8, pady=8)
        self.tab_control = ttk.Frame(self.nb)
        self.tab_reader = ttk.Frame(self.nb)
        self.tab_favs = ttk.Frame(self.nb)
        self.tab_kw = ttk.Frame(self.nb)
        self.tab_analytics = ttk.Frame(self.nb)

        self.nb.add(self.tab_control, text="🛰️ Control Panel")
        self.nb.add(self.tab_reader, text="📋 Intelligence Reader")
        self.nb.add(self.tab_favs, text="⭐ Favorites Hub")
        self.nb.add(self.tab_kw, text="🧠 Keyword Brain")
        self.nb.add(self.tab_analytics, text="📈 Policy Analytics")

        self._build_control_tab()
        self._build_reader_tab()
        self._build_favs_tab()
        self._build_kw_tab()
        self._build_analytics_tab()
        self.refresh_fav_hub()

    # ======================================================================
    # Control tab
    # ======================================================================
    def _build_control_tab(self) -> None:
        f = ttk.Frame(self.tab_control, padding=15)
        f.pack(fill="both", expand=True)
        f.columnconfigure(1, weight=1)
        f.rowconfigure(10, weight=1)

        ttk.Label(f, text="GOV.UK Advanced Intelligence Platform", font=("Segoe UI", 16, "bold"),
                  foreground="#00247D").grid(row=0, column=0, columnspan=3, pady=10)

        ttk.Label(f, text="Topic Query:").grid(row=1, column=0, sticky="w", pady=4)
        self.e_topic = ttk.Entry(f, width=40)
        self.e_topic.grid(row=1, column=1, sticky="ew", pady=4, padx=5)
        self.e_topic.insert(0, "blueprint modern digital government")
        ttk.Button(f, text="⭐ Save Fav Topic", command=self.add_favorite_topic).grid(row=1, column=2, padx=5)

        self.exact_match_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(f, text="Exact Phrase Match Only", variable=self.exact_match_var).grid(
            row=2, column=1, sticky="w", padx=5)

        ttk.Label(f, text="Target Count:").grid(row=3, column=0, sticky="w", pady=4)
        self.sp_count = ttk.Spinbox(f, from_=10, to=200, increment=10, width=10)
        self.sp_count.grid(row=3, column=1, sticky="w", pady=4, padx=5)
        self.sp_count.set(20)

        ttk.Label(f, text="Sort By:").grid(row=4, column=0, sticky="w", pady=4)
        self.cb_sort = ttk.Combobox(f, values=["Best Match", "Most Recent"], state="readonly", width=15)
        self.cb_sort.grid(row=4, column=1, sticky="w", pady=4, padx=5)
        self.cb_sort.set("Best Match")

        ttk.Label(f, text="Department:").grid(row=5, column=0, sticky="w", pady=4)
        self.cb_dept = ttk.Combobox(f, values=["All Departments", *config.DEPARTMENT_SLUGS.keys()],
                                    state="readonly", width=25)
        self.cb_dept.grid(row=5, column=1, sticky="w", pady=4, padx=5)
        self.cb_dept.set("All Departments")

        ttk.Label(f, text="Document Type:").grid(row=6, column=0, sticky="w", pady=4)
        self.cb_doc_type = ttk.Combobox(f, values=["All Types", *config.DOC_TYPE_SLUGS.keys()],
                                        state="readonly", width=25)
        self.cb_doc_type.grid(row=6, column=1, sticky="w", pady=4, padx=5)
        self.cb_doc_type.set("All Types")

        fav_frame = ttk.LabelFrame(f, text=" ⭐ Permanent Topic Favorites (Click to Load) ", padding=8)
        fav_frame.grid(row=7, column=0, columnspan=3, sticky="ew", pady=8)
        self.fav_buttons_inner = ttk.Frame(fav_frame)
        self.fav_buttons_inner.pack(fill="x")
        self.render_favorite_topics()

        btn_frame = ttk.Frame(f)
        btn_frame.grid(row=8, column=0, columnspan=3, pady=10)
        ttk.Button(btn_frame, text="🚀 Run Deep Search", command=self.start_harvest).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="⚔️ Split-Screen Compare", command=self.open_split_comparator).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="🔍 Cross-PDF Search", command=self.cross_pdf_search_dialog).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="📜 Search History", command=self.open_history_dialog).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="📄 Export Word Brief (.docx)", command=self.export_word_briefing).pack(side="left", padx=4)

        self.log_box = tk.Text(f, height=6, bg="#1e293b", fg="#f8fafc", font=("Consolas", 10))
        self.log_box.grid(row=10, column=0, columnspan=3, sticky="nsew", pady=10)

    def render_favorite_topics(self) -> None:
        for widget in self.fav_buttons_inner.winfo_children():
            widget.destroy()
        if not self.state.favorite_topics:
            ttk.Label(self.fav_buttons_inner, text="No favorite search terms saved yet.",
                      font=("Segoe UI", 9, "italic")).pack()
            return
        for fav in self.state.favorite_topics:
            btn_f = ttk.Frame(self.fav_buttons_inner)
            btn_f.pack(side="left", padx=4, pady=2)
            ttk.Button(btn_f, text=f"⭐ {fav}", command=lambda t=fav: self.load_fav_topic_search(t)).pack(side="left")
            ttk.Button(btn_f, text="❌", width=2, command=lambda t=fav: self.remove_favorite_topic(t)).pack(side="left", padx=2)

    def add_favorite_topic(self) -> None:
        if self.state.add_favorite_topic(self.e_topic.get().strip()):
            self.render_favorite_topics()

    def remove_favorite_topic(self, topic: str) -> None:
        if self.state.remove_favorite_topic(topic):
            self.render_favorite_topics()

    def load_fav_topic_search(self, topic: str) -> None:
        self.e_topic.delete(0, tk.END)
        self.e_topic.insert(0, topic)
        self.start_harvest()

    # -- logging (thread-safe) ------------------------------------------------
    def log(self, msg: str) -> None:
        """Safe to call from any thread: schedules the actual write on the main thread."""
        self.root.after(0, lambda: self._log_on_main_thread(msg))

    def _log_on_main_thread(self, msg: str) -> None:
        self.log_box.insert(tk.END, msg + "\n")
        self.log_box.see(tk.END)

    # -- harvesting ------------------------------------------------------------
    def start_harvest(self) -> None:
        topic = self.e_topic.get().strip()
        if not topic:
            return
        self.state.record_search(topic, self.cb_dept.get(), self.cb_doc_type.get())
        threading.Thread(target=self._run_harvest, args=(topic,), daemon=True).start()

    def _run_harvest(self, topic: str) -> None:
        """Runs entirely on a background thread. Must not touch widgets directly."""
        self.log(f"🔎 Scanning GOV.UK for '{topic}'...")
        try:
            target_total = int(self.sp_count.get())
        except (tk.TclError, ValueError):
            target_total = 20

        raw_results = deep_harvest_gov_uk(
            topic, target_total, self.cb_sort.get(), self.cb_dept.get(),
            self.cb_doc_type.get(), self.exact_match_var.get(), log_cb=self.log,
        )
        topic_dir, docs, _suggestions = archive.build_and_save_archive(topic, raw_results, log_cb=self.log)
        if docs:
            self.log(f"✅ Saved {len(docs)} documents for '{topic}'.")
            self.root.after(0, lambda: self._sync_reader(topic_dir, docs))

    def _sync_reader(self, topic_dir: Path, docs: list[Document]) -> None:
        self.active_docs = docs
        briefing_text = archive.load_briefing_text(topic_dir)
        self.txt_briefing.delete("1.0", tk.END)
        self.txt_briefing.insert(tk.END, briefing_text)
        self.refresh_doc_list()
        self.generate_analytics_matrix()
        self.nb.select(self.tab_reader)

    # ======================================================================
    # Reader tab
    # ======================================================================
    def _build_reader_tab(self) -> None:
        paned = ttk.Panedwindow(self.tab_reader, orient="horizontal")
        paned.pack(fill="both", expand=True, padx=5, pady=5)

        left = ttk.LabelFrame(paned, text=" Discovered Sources ", padding=8)
        paned.add(left, weight=1)

        self.e_filter = ttk.Entry(left)
        self.e_filter.pack(fill="x", pady=4)
        self.e_filter.bind("<KeyRelease>", lambda e: self.refresh_doc_list())

        self.lb_docs = tk.Listbox(left, bg="#ffffff", selectbackground="#00247D", exportselection=False)
        self.lb_docs.pack(fill="both", expand=True)
        self.lb_docs.bind("<<ListboxSelect>>", self.on_doc_select)
        self.lb_docs.bind("<Double-1>", self.on_doc_double_click)

        # Vertical PanedWindow on the right to make Details vs Viewer draggable
        right_paned = ttk.Panedwindow(paned, orient="vertical")
        paned.add(right_paned, weight=3)

        right_top = ttk.LabelFrame(right_paned, text=" Document Details ", padding=8)
        right_paned.add(right_top, weight=1)

        btn_row = ttk.Frame(right_top)
        btn_row.pack(fill="x", pady=(0, 4))
        self.btn_star = ttk.Button(btn_row, text="⭐ Star Source", command=self.toggle_starred)
        self.btn_star.pack(side="left", padx=2)
        ttk.Button(btn_row, text="🏷️ Set Tag", command=self.apply_custom_doc_tag).pack(side="left", padx=2)
        ttk.Button(btn_row, text="📋 Copy Citation", command=self.copy_citation_popup).pack(side="left", padx=2)

        self.tab_details = ttk.Frame(right_top)
        self.tab_details.pack(fill="both", expand=True)

        self.txt_details = scrolledtext.ScrolledText(self.tab_details, height=8, wrap="word")
        self.txt_details.pack(fill="both", expand=True, padx=2, pady=2)

        ttk.Label(self.tab_details, text="Attachments (double-click to open):").pack(anchor="w", padx=5)

        # Attachment listbox with dedicated scrollbar
        att_frame = ttk.Frame(self.tab_details)
        att_frame.pack(fill="both", expand=True, padx=5, pady=2)

        sb_att = ttk.Scrollbar(att_frame, orient="vertical")
        self.lb_atts = tk.Listbox(att_frame, height=4, exportselection=False, yscrollcommand=sb_att.set)
        sb_att.config(command=self.lb_atts.yview)

        sb_att.pack(side="right", fill="y")
        self.lb_atts.pack(side="left", fill="both", expand=True)
        self.lb_atts.bind("<Double-1>", self.on_attachment_open)

        right_bottom = ttk.Frame(right_paned)
        right_paned.add(right_bottom, weight=2)

        self.reader_nb = ttk.Notebook(right_bottom)
        self.reader_nb.pack(fill="both", expand=True)
        self.reader_pdf_viewer = PDFViewerWidget(self.reader_nb, lambda: self.state.keyword_rules)
        self.reader_nb.add(self.reader_pdf_viewer, text="📕 PDF Viewer")

        briefing_frame = ttk.Frame(self.reader_nb)
        self.reader_nb.add(briefing_frame, text="📰 Briefing")
        self.txt_briefing = scrolledtext.ScrolledText(briefing_frame, wrap="word")
        self.txt_briefing.pack(fill="both", expand=True)

    def refresh_doc_list(self) -> None:
        self.lb_docs.delete(0, tk.END)
        q = self.e_filter.get().strip().lower()
        for doc in self.active_docs:
            if q and q not in doc.title.lower() and q not in doc.description.lower():
                continue
            icon = "⭐" if self.state.is_favorite(doc) else "📄"
            tag = self.state.get_tag(doc)
            tag_label = f" [{tag}]" if tag else ""
            self.lb_docs.insert(tk.END, f"{icon} {doc.title}{tag_label}")

    def _visible_docs(self) -> list[Document]:
        """Docs currently shown in lb_docs, honoring the active filter (same order)."""
        q = self.e_filter.get().strip().lower()
        if not q:
            return self.active_docs
        return [d for d in self.active_docs if q in d.title.lower() or q in d.description.lower()]

    def on_doc_select(self, _event) -> None:
        sel = self.lb_docs.curselection()
        if not sel:
            return
        docs = self._visible_docs()
        if sel[0] >= len(docs):
            return
        doc = docs[sel[0]]
        self.selected_doc = doc

        details = (
            f"SOURCE: {doc.title}\nDATE: {doc.date}\nURL: {doc.url}\n"
            f"{'-' * 40}\n{doc.description}"
        )
        self.txt_details.delete("1.0", tk.END)
        self.txt_details.insert(tk.END, details)

        self.btn_star.config(text="🌟 Unstar Source" if self.state.is_favorite(doc) else "⭐ Star Source")

        self.active_atts = doc.attachments
        self.lb_atts.delete(0, tk.END)
        for url in self.active_atts:
            self.lb_atts.insert(tk.END, _attachment_label(url))

    def on_doc_double_click(self, _event) -> None:
        if self.selected_doc:
            webbrowser.open_new_tab(self.selected_doc.url)

    def toggle_starred(self) -> None:
        if not self.selected_doc:
            return
        is_fav = self.state.toggle_favorite_source(self.selected_doc)
        self.btn_star.config(text="🌟 Unstar Source" if is_fav else "⭐ Star Source")
        self.refresh_doc_list()
        self.refresh_fav_hub()

    def apply_custom_doc_tag(self) -> None:
        if not self.selected_doc:
            return
        tag = simpledialog.askstring("Set Custom Tag", "Enter custom label for this document (e.g., 'High Priority', 'Legal'):")
        if tag:
            self.state.set_tag(self.selected_doc, tag)
            self.refresh_doc_list()

    def on_attachment_open(self, _event) -> None:
        sel = self.lb_atts.curselection()
        if sel and self.active_atts:
            self._route_attachment_open(self.active_atts[sel[0]], viewer=self.reader_pdf_viewer, notebook=self.reader_nb)

    # -- attachment routing (PDF download now runs off the main thread) --------
    def _route_attachment_open(self, url: str, viewer: PDFViewerWidget, notebook: ttk.Notebook) -> None:
        tag = classify_attachment_url(url)
        if "📕 PDF" in tag:
            self._open_pdf_attachment(url, viewer, notebook)
        elif "📊 Data" in tag:
            DataGridViewerWindow(self.root, url, title=url.rsplit("/", 1)[-1])
        else:
            webbrowser.open_new_tab(url)

    def _open_pdf_attachment(self, url: str, viewer: PDFViewerWidget, notebook: ttk.Notebook) -> None:
        pdf_name = sanitize_filename(url)
        local_p = os.path.abspath(os.path.join(config.DATA_DIR, pdf_name))

        if os.path.exists(local_p):
            viewer.load_pdf(local_p)
            notebook.select(viewer)
            return

        self.log(f"⬇️ Downloading {pdf_name}...")

        def _download():
            try:
                res = requests.get(url, timeout=30)
                res.raise_for_status()
            except requests.RequestException as exc:
                logger.warning("Failed to download PDF %s: %s", url, exc)
                self.root.after(0, lambda: messagebox.showerror("Download Failed", f"Could not download PDF:\n{exc}"))
                return
            try:
                os.makedirs(os.path.dirname(local_p), exist_ok=True)
                with open(local_p, "wb") as f:
                    f.write(res.content)
            except OSError as exc:
                self.root.after(0, lambda: messagebox.showerror("Save Failed", f"Could not save PDF locally:\n{exc}"))
                return
            self.root.after(0, lambda: (viewer.load_pdf(local_p), notebook.select(viewer)))

        threading.Thread(target=_download, daemon=True).start()

    # ======================================================================
    # Favorites Hub tab
    # ======================================================================
    def _build_favs_tab(self) -> None:
        paned = ttk.Panedwindow(self.tab_favs, orient="horizontal")
        paned.pack(fill="both", expand=True, padx=5, pady=5)

        left = ttk.LabelFrame(paned, text=" ⭐ Starred Sources ", padding=8)
        paned.add(left, weight=1)
        self.lb_favs = tk.Listbox(left, exportselection=False)
        self.lb_favs.pack(fill="both", expand=True)
        self.lb_favs.bind("<<ListboxSelect>>", self.on_fav_select)
        self.lb_favs.bind("<Double-1>", self.on_fav_double_click)

        # Vertical PanedWindow on the right for drag-resizing
        right_paned = ttk.Panedwindow(paned, orient="vertical")
        paned.add(right_paned, weight=3)

        right_top = ttk.LabelFrame(right_paned, text=" Details & Attachments ", padding=8)
        right_paned.add(right_top, weight=1)

        tab_det = ttk.Frame(right_top)
        tab_det.pack(fill="both", expand=True)

        self.txt_fav_details = scrolledtext.ScrolledText(tab_det, height=6, wrap="word")
        self.txt_fav_details.pack(fill="both", expand=True, padx=2, pady=2)

        ttk.Label(tab_det, text="Attachments for selected source:").pack(anchor="w", padx=5)

        # Scrollable listbox for favorite attachments
        fav_att_frame = ttk.Frame(tab_det)
        fav_att_frame.pack(fill="both", expand=True, padx=5, pady=2)

        sb_fav_att = ttk.Scrollbar(fav_att_frame, orient="vertical")
        self.lb_fav_atts = tk.Listbox(fav_att_frame, height=4, exportselection=False, yscrollcommand=sb_fav_att.set)
        sb_fav_att.config(command=self.lb_fav_atts.yview)

        sb_fav_att.pack(side="right", fill="y")
        self.lb_fav_atts.pack(side="left", fill="both", expand=True)
        self.lb_fav_atts.bind("<Double-1>", self.on_fav_att_open)

        att_nb = ttk.Notebook(tab_det)
        att_nb.pack(fill="both", expand=True, pady=(6, 0))

        pdf_tab = ttk.Frame(att_nb)
        att_nb.add(pdf_tab, text="📕 All PDFs")
        self.lb_fav_pdfs = tk.Listbox(pdf_tab, exportselection=False)
        self.lb_fav_pdfs.pack(fill="both", expand=True)
        self.lb_fav_pdfs.bind("<Double-1>", lambda e: self._open_filtered_fav_asset(self.lb_fav_pdfs, self.all_fav_attachments["pdf"]))

        link_tab = ttk.Frame(att_nb)
        att_nb.add(link_tab, text="🌐 All Links")
        self.lb_fav_links = tk.Listbox(link_tab, exportselection=False)
        self.lb_fav_links.pack(fill="both", expand=True)
        self.lb_fav_links.bind("<Double-1>", lambda e: self._open_filtered_fav_asset(self.lb_fav_links, self.all_fav_attachments["link"]))

        data_tab = ttk.Frame(att_nb)
        att_nb.add(data_tab, text="📊 All Datasets")
        self.lb_fav_data = tk.Listbox(data_tab, exportselection=False)
        self.lb_fav_data.pack(fill="both", expand=True)
        self.lb_fav_data.bind("<Double-1>", lambda e: self._open_filtered_fav_asset(self.lb_fav_data, self.all_fav_attachments["data"]))

        right_bottom = ttk.Frame(right_paned)
        right_paned.add(right_bottom, weight=2)

        self.fav_nb = ttk.Notebook(right_bottom)
        self.fav_nb.pack(fill="both", expand=True)
        self.fav_pdf_viewer = PDFViewerWidget(self.fav_nb, lambda: self.state.keyword_rules)
        self.fav_nb.add(self.fav_pdf_viewer, text="📕 Preview")

    def _favorite_ids(self) -> list[str]:
        return list(self.state.favorite_sources.keys())

    def refresh_fav_hub(self) -> None:
        self.lb_favs.delete(0, tk.END)
        self.lb_fav_pdfs.delete(0, tk.END)
        self.lb_fav_links.delete(0, tk.END)
        self.lb_fav_data.delete(0, tk.END)
        self.all_fav_attachments = {"pdf": [], "link": [], "data": []}

        for doc_id in self._favorite_ids():
            entry = self.state.favorite_sources[doc_id]
            self.lb_favs.insert(tk.END, f"⭐ {entry['title']}")
            for url in entry.get("attachments", []):
                tag = classify_attachment_url(url)
                item_str = f"{_attachment_label(url)}  [{entry['title'][:25]}...]"
                bucket = "pdf" if "📕 PDF" in tag else ("data" if "📊 Data" in tag else "link")
                self.all_fav_attachments[bucket].append(url)
                target_lb = {"pdf": self.lb_fav_pdfs, "data": self.lb_fav_data, "link": self.lb_fav_links}[bucket]
                target_lb.insert(tk.END, item_str)

    def on_fav_select(self, _event) -> None:
        ids = self._favorite_ids()
        sel = self.lb_favs.curselection()
        if not sel:
            return
        doc_id = ids[sel[0]]
        entry = self.state.favorite_sources[doc_id]
        self.fav_selected_doc_id = doc_id
        self.fav_active_atts = list(entry.get("attachments", []))

        self.txt_fav_details.delete("1.0", tk.END)
        self.txt_fav_details.insert(
            tk.END, f"TITLE: {entry['title']}\nURL: {entry['url']}\nTOPIC: {entry.get('topic', '')}"
        )

        self.lb_fav_atts.delete(0, tk.END)
        for url in self.fav_active_atts:
            self.lb_fav_atts.insert(tk.END, _attachment_label(url))

    def on_fav_double_click(self, _event) -> None:
        ids = self._favorite_ids()
        sel = self.lb_favs.curselection()
        if sel:
            webbrowser.open_new_tab(self.state.favorite_sources[ids[sel[0]]]["url"])

    def _open_filtered_fav_asset(self, listbox_widget: tk.Listbox, asset_url_list: list[str]) -> None:
        sel = listbox_widget.curselection()
        if sel and asset_url_list:
            self._route_attachment_open(asset_url_list[sel[0]], viewer=self.fav_pdf_viewer, notebook=self.fav_nb)

    def on_fav_att_open(self, _event) -> None:
        sel = self.lb_fav_atts.curselection()
        if sel and self.fav_active_atts:
            self._route_attachment_open(self.fav_active_atts[sel[0]], viewer=self.fav_pdf_viewer, notebook=self.fav_nb)

    # ======================================================================
    # Keyword Brain tab
    # ======================================================================
    def _build_kw_tab(self) -> None:
        f = ttk.Frame(self.tab_kw, padding=10)
        f.pack(fill="both", expand=True)

        left = ttk.LabelFrame(f, text=" Categories ", padding=8)
        left.pack(side="left", fill="y", padx=(0, 8))
        self.lb_kw_cats = tk.Listbox(left, exportselection=False, width=28)
        self.lb_kw_cats.pack(fill="y", expand=True)
        self.lb_kw_cats.bind("<<ListboxSelect>>", self.on_kw_cat_select)

        cat_row = ttk.Frame(left)
        cat_row.pack(fill="x", pady=4)
        self.e_new_cat = ttk.Entry(cat_row)
        self.e_new_cat.pack(side="left", fill="x", expand=True)
        ttk.Button(cat_row, text="➕", width=3, command=self.add_kw_category).pack(side="left")
        ttk.Button(cat_row, text="🗑️", width=3, command=self.remove_kw_category).pack(side="left")

        right = ttk.LabelFrame(f, text=" Terms ", padding=8)
        right.pack(side="left", fill="both", expand=True)
        self.lb_kw_terms = tk.Listbox(right, exportselection=False)
        self.lb_kw_terms.pack(fill="both", expand=True)
        self.lb_kw_terms.bind("<Double-1>", self.toggle_kw_term_state)

        term_row = ttk.Frame(right)
        term_row.pack(fill="x", pady=4)
        self.e_new_kw = ttk.Entry(term_row)
        self.e_new_kw.pack(side="left", fill="x", expand=True)
        ttk.Button(term_row, text="➕ Add Term", command=self.add_kw_term).pack(side="left", padx=2)
        ttk.Button(term_row, text="🗑️ Remove Term", command=self.remove_kw_term).pack(side="left", padx=2)

        self.refresh_kw_categories_list()

    def refresh_kw_categories_list(self) -> None:
        self.lb_kw_cats.delete(0, tk.END)
        for cat in self.state.keyword_rules:
            self.lb_kw_cats.insert(tk.END, cat)
        if self.state.keyword_rules:
            self.lb_kw_cats.select_set(0)
            self.on_kw_cat_select(None)

    def add_kw_category(self) -> None:
        if self.state.add_keyword_category(self.e_new_cat.get()):
            self.e_new_cat.delete(0, tk.END)
            self.refresh_kw_categories_list()

    def remove_kw_category(self) -> None:
        sel = self.lb_kw_cats.curselection()
        if not sel:
            return
        cat = self.lb_kw_cats.get(sel[0])
        if self.state.remove_keyword_category(cat):
            self.refresh_kw_categories_list()
            self.lb_kw_terms.delete(0, tk.END)

    def on_kw_cat_select(self, _event) -> None:
        sel = self.lb_kw_cats.curselection()
        if not sel:
            return
        cat = self.lb_kw_cats.get(sel[0])
        self.lb_kw_terms.delete(0, tk.END)
        for term, enabled in self.state.keyword_rules.get(cat, {}).get("terms", {}).items():
            self.lb_kw_terms.insert(tk.END, f"{'☑' if enabled else '☐'} {term}")

    def add_kw_term(self) -> None:
        sel = self.lb_kw_cats.curselection()
        if not sel:
            return
        cat = self.lb_kw_cats.get(sel[0])
        if self.state.add_keyword_term(cat, self.e_new_kw.get()):
            self.on_kw_cat_select(None)
            self.e_new_kw.delete(0, tk.END)

    def toggle_kw_term_state(self, event=None) -> None:
        cat_sel = self.lb_kw_cats.curselection()
        if not cat_sel:
            return
        index = self.lb_kw_terms.nearest(event.y) if (event and hasattr(event, "y")) else None
        if index is None:
            term_sel = self.lb_kw_terms.curselection()
            if not term_sel:
                return
            index = term_sel[0]
        cat = self.lb_kw_cats.get(cat_sel[0])
        item_str = self.lb_kw_terms.get(index)
        if not item_str:
            return
        term = item_str[2:].strip()
        self.state.toggle_keyword_term(cat, term)
        self.on_kw_cat_select(None)
        self.lb_kw_terms.select_set(index)

    def remove_kw_term(self) -> None:
        cat_sel, term_sel = self.lb_kw_cats.curselection(), self.lb_kw_terms.curselection()
        if not (cat_sel and term_sel):
            return
        cat = self.lb_kw_cats.get(cat_sel[0])
        term = self.lb_kw_terms.get(term_sel[0])[2:].strip()
        if self.state.remove_keyword_term(cat, term):
            self.on_kw_cat_select(None)

    # ======================================================================
    # Analytics tab
    # ======================================================================
    def _build_analytics_tab(self) -> None:
        f = ttk.Frame(self.tab_analytics, padding=10)
        f.pack(fill="both", expand=True)
        columns = ["title"] + list(self.state.keyword_rules.keys())
        self.analytics_tree = ttk.Treeview(f, columns=columns, show="headings")
        for col in columns:
            self.analytics_tree.heading(col, text=col)
            self.analytics_tree.column(col, width=140 if col != "title" else 320, anchor="w")
        self.analytics_tree.pack(fill="both", expand=True)
        self.analytics_tree.bind("<Double-1>", self.on_analytics_doc_double_click)

    def generate_analytics_matrix(self) -> None:
        for row in self.analytics_tree.get_children():
            self.analytics_tree.delete(row)
        categories = list(self.state.keyword_rules.keys())
        for idx, doc in enumerate(self.active_docs):
            haystack = f"{doc.title} {doc.description}".lower()
            row_vals = [doc.title]
            for cat in categories:
                terms = self.state.keyword_rules[cat].get("terms", {})
                active_terms = [t for t, enabled in terms.items() if enabled]
                matches = sum(haystack.count(t) for t in active_terms)
                row_vals.append(matches)
            self.analytics_tree.insert("", "end", iid=str(idx), values=row_vals)

    def on_analytics_doc_double_click(self, _event) -> None:
        sel = self.analytics_tree.selection()
        if not sel:
            return
        doc_idx = int(sel[0])
        if not (0 <= doc_idx < len(self.active_docs)):
            return
        self.lb_docs.selection_clear(0, tk.END)
        self.lb_docs.selection_set(doc_idx)
        self.lb_docs.see(doc_idx)
        self.on_doc_select(None)
        self.nb.select(self.tab_reader)

    # ======================================================================
    # Dialogs / misc tools
    # ======================================================================
    def open_split_comparator(self) -> None:
        SplitPDFComparatorWindow(self.root, lambda: self.state.keyword_rules)

    def open_history_dialog(self) -> None:
        pop = tk.Toplevel(self.root)
        pop.title("📜 Search History")
        pop.geometry("600x400")
        lb = tk.Listbox(pop, bg="#ffffff", font=("Segoe UI", 10))
        lb.pack(fill="both", expand=True, padx=10, pady=10)
        for rec in self.state.search_history:
            lb.insert(tk.END, f"[{rec.get('timestamp')}] Query: '{rec.get('query')}' | Dept: {rec.get('dept')}")

        def rerun(_event):
            sel = lb.curselection()
            if sel:
                self.e_topic.delete(0, tk.END)
                self.e_topic.insert(0, self.state.search_history[sel[0]].get("query"))
                pop.destroy()
                self.start_harvest()

        lb.bind("<Double-1>", rerun)

    def copy_citation_popup(self) -> None:
        if not self.selected_doc:
            return
        year = datetime.now().strftime("%Y")
        apa = f"HM Government. ({year}). {self.selected_doc.title}. GOV.UK."
        oscola = f"HM Government, '{self.selected_doc.title}' ({year}) accessed {datetime.now().strftime('%B %Y')}."
        pop = tk.Toplevel(self.root)
        pop.title("📋 Reference Generator")
        pop.geometry("550x250")
        ttk.Label(pop, text="APA Format:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=10, pady=(10, 2))
        t1 = tk.Text(pop, height=2, wrap="word")
        t1.pack(fill="x", padx=10)
        t1.insert(tk.END, apa)
        ttk.Label(pop, text="OSCOLA Format:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=10, pady=(10, 2))
        t2 = tk.Text(pop, height=2, wrap="word")
        t2.pack(fill="x", padx=10)
        t2.insert(tk.END, oscola)

    def export_word_briefing(self) -> None:
        if not HAS_DOCX:
            messagebox.showwarning("Library Missing", "Run 'pip install python-docx' first.")
            return
        briefing_text = self.txt_briefing.get("1.0", tk.END).strip()
        if not briefing_text:
            messagebox.showinfo("No Briefing", "Please run a search scan first.")
            return
        doc = docx.Document()
        doc.add_heading(f"GOV.UK Briefing: {self.e_topic.get().strip().upper()}", 0)
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(briefing_text)
        out_p = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if out_p:
            doc.save(out_p)
            messagebox.showinfo("Export Complete", f"Saved to: {out_p}")

    def cross_pdf_search_dialog(self) -> None:
        pop = tk.Toplevel(self.root)
        pop.title("🔍 Multi-PDF Cross Search")
        pop.geometry("600x450")
        e = ttk.Entry(pop)
        e.pack(fill="x", padx=10, pady=5)
        txt = scrolledtext.ScrolledText(pop, wrap="word")
        txt.pack(fill="both", expand=True, padx=10, pady=5)
        status = ttk.Label(pop, text="")
        status.pack(fill="x", padx=10)

        def run():
            term = e.get().strip().lower()
            if not term:
                return
            txt.delete("1.0", tk.END)
            status.config(text="Searching…")
            threading.Thread(target=self._cross_pdf_search_worker, args=(term, txt, status), daemon=True).start()

        ttk.Button(pop, text="Search PDFs", command=run).pack(pady=5)

    def _cross_pdf_search_worker(self, term: str, txt: scrolledtext.ScrolledText, status: ttk.Label) -> None:
        """Runs on a background thread; only schedules widget updates via `after`."""
        import fitz  # local import: keeps the dependency optional at module load time

        matches_found = 0
        for r, _dirs, files in os.walk(config.DATA_DIR):
            for file in files:
                if not file.endswith(".pdf"):
                    continue
                try:
                    doc = fitz.open(os.path.join(r, file))
                except (RuntimeError, OSError) as exc:
                    logger.info("Skipping unreadable PDF %s: %s", file, exc)
                    continue
                for p_num, page in enumerate(doc):
                    if term in page.get_text("text").lower():
                        matches_found += 1
                        line = f"📌 [File: {file} | Page {p_num + 1}]\n"
                        self.root.after(0, lambda line=line: txt.insert(tk.END, line))
        self.root.after(0, lambda: status.config(text=f"Done — {matches_found} match(es) found."))